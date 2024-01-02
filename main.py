from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from datetime import datetime
import time

# Obter a data e hora atual
agora = datetime.now()
data_atual = agora.strftime("%Y-%m-%d")
hora_atual = agora.strftime("%H-%M-%S")

# Caminho para o diretório desejado
diretorio_base = r' ~/Downloads/Otimiza-o2-main/Otimiza-o2-main'

# Criar o caminho para a pasta do dia
caminho = os.path.join(diretorio_base, data_atual)
if not os.path.exists(caminho):
    os.makedirs(caminho)

# Crie uma instância do navegador
driver = webdriver.Edge()

# Maximize a janela do navegador
driver.maximize_window()

# Abra a página web
driver.get("site da verificação")

# Aguarde até que os elementos estejam presentes na página
wait = WebDriverWait(driver, 10)

email_input = wait.until(EC.presence_of_element_located((By.NAME, "email")))
senha_input = wait.until(EC.presence_of_element_located((By.NAME, "pass")))
submit_button = wait.until(EC.element_to_be_clickable((By.ID, "submitbutton")))

email = "usuario"  # Insira seu email aqui
senha = "senha"  # Insira sua senha aqui

if email_input and senha_input:
    email_input.send_keys(email)
    senha_input.send_keys(senha)
    submit_button.click()

    

    # Aguarde o carregamento após o login (adicione tempo extra se necessário)
    time.sleep(5)  # Altere o tempo de espera conforme necessário para que a página carregue completamente

    # Capture a tela após o login na primeira aba
    nome_arquivo_1 = f"{data_atual}_{hora_atual}_Imagem_apos_login_aba1.png"
    screenshot_aba1 = driver.save_screenshot(os.path.join(caminho, nome_arquivo_1))

    # Abra a segunda página em uma nova aba
    driver.execute_script("window.open('outra página para tirar print');")

    # Aguarde o carregamento da segunda página
    time.sleep(2)  # Tempo para garantir o carregamento da segunda página

    # Troque para a segunda aba
    driver.switch_to.window(driver.window_handles[1])

    # Captura do segundo print
    nome_arquivo_2 = f"{data_atual}_{hora_atual}_Segunda_pagina.png"
    screenshot_2 = driver.save_screenshot(os.path.join(caminho, nome_arquivo_2))

    # Feche o navegador
    driver.quit()

    # Criar um documento do Word
    doc = Document()

    # Estilo do título com a data
    titulo = doc.add_heading(f'Gravação dos Ramais Operacionais', level=1)
    titulo = doc.add_heading(f'Data da Evidência: {data_atual}', level=1)
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Parágrafo para cada imagem
    for arquivo in [nome_arquivo_1, nome_arquivo_2]:
        paragrafo = doc.add_paragraph()
        paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        imagem_path = os.path.join(caminho, arquivo)
        doc.add_picture(imagem_path, width=Inches(5))  # Ajuste a largura conforme necessário

    # Salvar o documento
    doc_path = os.path.join(caminho, f'prints_{data_atual}.docx')
    doc.save(doc_path)
